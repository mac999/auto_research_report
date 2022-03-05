# Purpose: Generate research & meeting report automatically from website or blog
# Date: 2022.3.1
# Author: mac999
# 

import logging
from urllib.parse import urljoin
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import argparse
import sys
import csv

logging.basicConfig(
    format='%(asctime)s %(levelname)s:%(message)s',
    level=logging.INFO)

class auto_report:
    _urls = []
    _output = ""
    _begin = 0
    _dist = 0
    _step = 0
    _header = []
    _body_title = []

    def download_url(self, url):
        return requests.get(url).text

    def searchGoogle(self, query):
        try:
            from googlesearch import search
        except ImportError:
            print("No module named 'google' found")

        for u in search(query, tld="co.in", num=10, stop=10, pause=2):
            print(u)
            self._urls.append(u)

    def get_linked_urls(self, url, html):
        soup = BeautifulSoup(html, 'html.parser')
        for link in soup.find_all('a'):
            path = link.get('href')
            if path and path.startswith('/'):
                path = urljoin(url, path)
            yield path

    def is_empty_content(self, txt, ratio = 0.7):
        space = txt.count(' ')    
        r = 1.0 - float(space) / float(len(txt))
        if r < ratio:
            return True
        return False

    def generate_body(self, html, begin = 2000, dist = 500, step = 2000, cnt = 1):
        print("\nResearch Note")
        soup = BeautifulSoup(html, 'html.parser')

        body = ""
        index = 0
        for b in soup.find_all('body'):
            t = b.text
            t = t.replace("\n", "")

            b = begin
            if b >= len(t):
                b = 0
            if dist <= 0:
                dist = len(t) - b - 1

            current = 0
            while (b + dist + step) < len(t):
                t2 = t[b: b + dist]
                b = b + dist + step

                if self.is_empty_content(t2):
                    continue

                if current >= len(self._body_title):
                    break
                t2 = self._body_title[current] + "\n" + t2 + "...\n\n"
                print(t2)
                body = body + t2
                print("begin = ", b)
                current = current + 1

            if index >= cnt:
                break
            index = index + 1
        return body

    def crawl(self, url):
        html = self.download_url(url)
        body = self.parsing_body(html)

    def createWord(self, fname = "out.docx", body = "Hello world!"):
        document = Document()
        document.add_heading('Research Note', 0)

        table = document.add_table(rows=3, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Department'
        hdr_cells[1].text = self._header[0]
        hdr_cells[2].text = 'Position'
        hdr_cells[3].text = self._header[1]
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = self._header[2]
        hdr_cells[2].text = 'Name'
        hdr_cells[3].text = self._header[3]
        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = self._date
        hdr_cells[2].text = 'Place'
        hdr_cells[3].text = self._header[4]

        table = document.add_table(rows=2, cols=1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Note'
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = body

        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Confirmation'
        hdr_cells[1].text = 'Position'
        hdr_cells[2].text = self._header[0]
        hdr_cells[3].text = 'Name'
        hdr_cells[4].text = self._header[3] + ' (Sign)'

        # document.add_page_break()
        document.save(fname)

    def set_parameters(self, argv):
        url = 'https://sites.google.com/site/bimprinciple/in-the-news/gieob-uidijiteoljeonhwanguhyeon'
        # url = 'http://daddynkidsmakers.blogspot.com/2022/02/floor-plan.html'        

        if len(argv) >= 7:
            self._date = argv[1]
            self._urls.append(argv[2])
            if argv[2].find('http') < 0:
                print("Knowledge searching mode...")
                self._urls.clear()
                self.searchGoogle(argv[2])

            self._output = argv[3]
            self._begin = int(argv[4])
            self._dist = int(argv[5])
            self._step = int(argv[6])

    def read_config(self, fname = "config.csv"):
        print("\nReading config file...")
        index = 0
        with open(fname, 'r', encoding = 'UTF-8') as csvfile:
            reader = csv.reader(csvfile, delimiter=',')
            for row in reader:
                if index == 0:
                    self._header = row
                if index == 1:
                    self._body_title = row 
                index = index + 1

    def gen_note(self, num = 1):
        print("\nProcessing...")
        print(self._header)        
        print(self._body_title)        

        index = 0
        for u in self._urls:
            if index >= num:
                break
            index = index + 1
            html = self.download_url(u)
            body = self.generate_body(html, self._begin, self._dist, self._step)
            self.createWord(self._output, body)

if __name__ == '__main__':
    c = auto_report()
    ''' for test
     argv = []
    argv.append("")
    argv.append("2022/1/4")
    argv.append("https://sites.google.com/site/bimprinciple/in-the-news/gieob-uidijiteoljeonhwanguhyeon")
    argv.append("r3.docx")
    argv.append("100")
    argv.append("300")
    argv.append("50")
    c.set_parameters(argv) '''
    c.set_parameters(sys.argv)
    c.read_config()
    c.gen_note()



